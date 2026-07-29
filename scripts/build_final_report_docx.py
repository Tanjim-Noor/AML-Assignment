"""Build the submission DOCX from the final reviewable Markdown.

The builder starts from the retained Word artifact so its styles/theme remain
available, then reconstructs the content with a small, explicit academic style
system. It deliberately avoids code-facing field names and direct formatting
drift in the visible document.
"""

from __future__ import annotations

import argparse
import io
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "Assignment Report/AML_Assignment_Report_Final.md"
REFERENCE_DOCX = (
    ROOT
    / "Assignment Report/_backup/docx-2026-07-29/AML_Assignment_Report_Updated_Fixed.docx"
)
OUTPUT_DOCX = ROOT / "Assignment Report/AML_Assignment_Report_Final_Readable_Figures.docx"
ASSETS = ROOT / "Assignment Report/assets"

DENSE_FIGURE_ASSETS = {
    "fig02a_context_numeric_distributions.png",
    "fig02b_ai_numeric_distributions.png",
    "fig02c_excluded_outcome_distributions.png",
    "fig03a_context_categorical_distributions.png",
    "fig03b_ai_categorical_distributions.png",
    "fig04_correlation_and_target_relationships.png",
    "fig06_gpa_change_eda.png",
    "fig07_model_test_rmse_and_actual_vs_predicted.png",
    "fig08_residual_diagnostics.png",
    "fig09_direction_specific_errors.png",
}

FIGURE_ALT_TEXT = {
    "fig01_data_quality_and_schema.png": (
        "Data-quality audit with zero missing, duplicate and invalid-range records, "
        "beside the original field-type composition."
    ),
    "fig02a_context_numeric_distributions.png": (
        "Three distributions for previous-semester GPA, traditional study hours "
        "and exam anxiety."
    ),
    "fig02b_ai_numeric_distributions.png": (
        "Three distributions for weekly generative-AI hours, tool diversity and "
        "perceived AI dependency."
    ),
    "fig02c_excluded_outcome_distributions.png": (
        "Distributions of post-semester GPA and skill retention, both identified "
        "as post-outcome fields excluded from prediction."
    ),
    "fig03a_context_categorical_distributions.png": (
        "Frequency bar charts for major category, year of study and institutional "
        "AI policy."
    ),
    "fig03b_ai_categorical_distributions.png": (
        "Frequency bar charts for primary AI use case, prompt-engineering skill "
        "and paid subscription."
    ),
    "fig04_correlation_and_target_relationships.png": (
        "Numeric correlation matrix above a scatter plot of previous-semester GPA "
        "against semester GPA change."
    ),
    "fig05_gpa_direction_imbalance.png": (
        "Observed GPA-change direction counts beside the retained continuous target "
        "distributions."
    ),
    "fig06_gpa_change_eda.png": (
        "Four-panel exploration of GPA-change distribution, previous GPA, prompt "
        "skill and weekly AI-hours quartiles."
    ),
    "fig07_model_test_rmse_and_actual_vs_predicted.png": (
        "Reserved-test RMSE comparison above actual-versus-predicted GPA change for "
        "the tuned histogram gradient boosting model."
    ),
    "fig08_residual_diagnostics.png": (
        "Residuals against predicted GPA change above the residual distribution for "
        "the tuned histogram gradient boosting model."
    ),
    "fig09_direction_specific_errors.png": (
        "Direction-specific MAE and RMSE above residual distributions for GPA "
        "decreases, unchanged values and increases."
    ),
    "fig10_permutation_importance.png": (
        "Horizontal bars ranking test-set permutation importance for the tuned "
        "histogram gradient boosting model."
    ),
}

FONT = "Times New Roman"
BODY_SIZE = 12
USABLE_WIDTH_DXA = 9030


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_spacing(paragraph, line=1.5, before=0, after=6):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_keep(paragraph, next_=False, lines=False):
    ppr = paragraph._p.get_or_add_pPr()
    if next_:
        ppr.append(OxmlElement("w:keepNext"))
    if lines:
        ppr.append(OxmlElement("w:keepLines"))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def shade_cell(cell, fill="E7E6E6"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=12)


def set_page_number_format(section, fmt, start=None):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    if start is not None:
        node.set(qn("w:start"), str(start))


def set_header(section, left, right):
    section.header.is_linked_to_previous = False
    for table in list(section.header.tables):
        table._element.getparent().remove(table._element)
    p = section.header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.27), WD_TAB_ALIGNMENT.RIGHT)
    set_spacing(p, line=1.0, after=0)
    r = p.add_run(left)
    set_run_font(r, size=10)
    r = p.add_run("\t" + right)
    set_run_font(r, size=10)


def set_footer(section, enabled=True):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.clear()
    if enabled:
        add_page_number(p)


def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(paragraph, text, allow_italics=False):
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("*") and allow_italics:
            set_run_font(paragraph.add_run(token[1:-1]), italic=True)
        else:
            set_run_font(paragraph.add_run(token.strip("*`")))
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def set_image_alt(picture_run, description):
    for node in picture_run._r.xpath(".//wp:docPr"):
        node.set("descr", description)
    for node in picture_run._r.xpath(".//pic:cNvPr"):
        node.set("name", description)


def compatible_image(path):
    """Return a stream that python-docx can reliably identify."""
    image = Image.open(path).convert("RGB")
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def mark_orphan_header_table_rows(docx_path):
    """Make inherited, unreferenced header tables harmless to a11y scanners."""
    temporary = docx_path.with_suffix(".a11y.tmp.docx")
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temporary, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if re.fullmatch(r"word/header\d+\.xml", info.filename):
                root = etree.fromstring(data)
                changed = False
                for table in root.xpath(".//w:tbl", namespaces=namespace):
                    rows = table.xpath("./w:tr", namespaces=namespace)
                    if not rows:
                        continue
                    row = rows[0]
                    tr_pr = row.find(qn("w:trPr"))
                    if tr_pr is None:
                        tr_pr = etree.Element(qn("w:trPr"))
                        row.insert(0, tr_pr)
                    if tr_pr.find(qn("w:tblHeader")) is None:
                        flag = etree.Element(qn("w:tblHeader"))
                        flag.set(qn("w:val"), "true")
                        tr_pr.append(flag)
                        changed = True
                if changed:
                    data = etree.tostring(
                        root, xml_declaration=True, encoding="UTF-8", standalone=True
                    )
            target.writestr(info, data)
    temporary.replace(docx_path)


def parse_markdown(md_text):
    body = md_text[md_text.index("# 1 Introduction") :]
    lines = body.splitlines()
    blocks = []
    i = 0
    paragraph_lines = []

    def flush():
        nonlocal paragraph_lines
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(x.strip() for x in paragraph_lines)))
            paragraph_lines = []

    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            flush()
            i += 1
            continue
        if line.startswith("# "):
            flush()
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            flush()
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        image = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image:
            flush()
            blocks.append(("image", (image.group(1), image.group(2))))
            i += 1
            continue
        if re.match(r"^Table \d+\.", line):
            flush()
            blocks.append(("table_caption", line))
            i += 1
            continue
        if re.match(r"^Figure \d+(?:\([a-z]\))?\.", line):
            flush()
            blocks.append(("figure_caption", line))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|\s*:?-+", lines[i + 1]):
            flush()
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue
        if re.match(r"^\d+\.\s+", line):
            flush()
            blocks.append(("numbered", re.sub(r"^\d+\.\s+", "", line)))
            i += 1
            continue
        if line.startswith("- "):
            flush()
            blocks.append(("reference" if any(b == ("h1", "References") for b in blocks) else "bullet", line[2:]))
            i += 1
            continue
        paragraph_lines.append(line)
        i += 1
    flush()
    return blocks


def get_front_text(md_text):
    title = re.search(r"\*\*(Predicting Semester GPA Change.+?)\*\*", md_text).group(1)
    abstract = re.search(r"\*\*Abstract\*\*\s*\n\n(.+?)\n\n<!-- Page break", md_text, re.S).group(1).strip()
    return title, abstract


def apply_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Heading 1", 15, 12, 8),
        ("Heading 2", 13, 10, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Number", "List Bullet"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(12)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(compatible_image(ASSETS / "apu-logo.jpg"), width=Inches(1.34))
    set_image_alt(run, "Asia Pacific University of Technology and Innovation logo")
    set_spacing(p, line=1.0, after=12)

    for text, size, after in (
        ("ASIA PACIFIC UNIVERSITY OF TECHNOLOGY & INNOVATION", 14, 10),
        ("Individual Assignment", 14, 18),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=size, bold=True)
        set_spacing(p, line=1.0, after=after)

    metadata = [
        ("Module code", "CT046-3-M"),
        ("Module name", "Applied Machine Learning"),
        ("Intake code", "CSSE / CT046-3-M-AML-L-1 / 2026-01-26"),
        ("Hand-out date", "25 May 2026"),
        ("Hand-in date", "24 July 2026"),
        ("Weightage", "100%"),
        ("Student name", "Tanjim Noor"),
        ("Student ID", "TP166089"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.6), WD_TAB_ALIGNMENT.LEFT)
        set_spacing(p, line=1.15, after=3)
        set_run_font(p.add_run(label), bold=True)
        set_run_font(p.add_run("\t" + value))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Declaration"), size=14, bold=True)
    set_spacing(p, line=1.0, before=12, after=8)
    for text in (
        "I declare that this work is my own except where the work of others is clearly acknowledged and cited. I have not presented it, in whole or in part, for another academic award.",
        "Generative-AI assistance was used for literature organisation, code review, language editing and document preparation. The student verified the sources, numerical results, interpretations and final wording and accepts responsibility for the submitted work.",
    ):
        p = doc.add_paragraph()
        add_inline(p, text)
        set_spacing(p, line=1.15, after=6)

    for label, value in (("Student name", "Tanjim Noor"), ("Student ID", "TP166089")):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.6), WD_TAB_ALIGNMENT.LEFT)
        set_spacing(p, line=1.15, after=3)
        set_run_font(p.add_run(label), bold=True)
        set_run_font(p.add_run("\t" + value))


def add_front_matter(doc, title, abstract, headings, figures, tables, page_map):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, line=1.15, before=72, after=20)
    set_run_font(p.add_run(title), size=18, bold=True)
    for text, bold, after in (
        ("CT046-3-M: Applied Machine Learning", True, 18),
        ("Submitted by", False, 10),
        ("Tanjim Noor", True, 3),
        ("TP166089", False, 12),
        ("Asia Pacific University of Technology & Innovation", False, 12),
        ("Submission date: 31 July 2026", False, 0),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), bold=bold)
        set_spacing(p, line=1.15, after=after)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Abstract"), size=15, bold=True)
    set_spacing(p, line=1.0, after=10)
    p = doc.add_paragraph()
    add_inline(p, abstract)
    set_spacing(p, line=1.5, after=6)

    def add_list_page(title_text, entries):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(title_text), size=15, bold=True)
        set_spacing(p, line=1.0, after=10)
        for label, level in entries:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.25 if level == 1 else 0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25 if level == 1 else -0.25)
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.05),
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS,
            )
            set_spacing(p, line=1.15, after=3)
            display_label = label
            if title_text == "List of Figures":
                first_sentence = re.match(r"^(Figure \d+\.\s+.*?)(?:\.\s|$)", label)
                if first_sentence:
                    display_label = first_sentence.group(1)
            set_run_font(p.add_run(display_label))
            page = page_map.get(label, "")
            if page:
                set_run_font(p.add_run("\t" + str(page)))

    toc_entries = []
    for kind, text in headings:
        toc_entries.append((text, 1 if kind == "h1" else 2))
    add_list_page("Table of Contents", toc_entries)
    add_list_page("List of Figures", [(x, 1) for x in figures])
    add_list_page("List of Tables", [(x, 1) for x in tables])


def table_widths(rows):
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        maximum = max(len(row[index]) if index < len(row) else 0 for row in rows)
        weights.append(max(8, min(maximum, 45)))
    minimum = 1200 if columns <= 4 else 900 if columns == 5 else 650
    remaining = USABLE_WIDTH_DXA - (minimum * columns)
    total = sum(weights)
    widths = [minimum + round(remaining * w / total) for w in weights]
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    return widths


def add_body(doc, blocks):
    in_references = False
    first_h1 = True
    last_image = None
    last_image_dense = False
    for kind, value in blocks:
        if kind == "h1":
            in_references = value == "References"
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, value)
            if value in {
                "2 Related Works",
                "3 Methods",
                "4 Dataset Preparation",
                "5 Model Implementation",
                "6 Model Validation",
                "References",
            }:
                p.paragraph_format.page_break_before = True
            else:
                p.paragraph_format.page_break_before = False
            first_h1 = False
            set_keep(p, next_=True, lines=True)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, value)
            set_keep(p, next_=True, lines=True)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            add_inline(p, value, allow_italics=in_references)
            set_spacing(p, line=1.5, after=6)
            set_keep(p, lines=True)
            if value.startswith("**Panel "):
                set_keep(p, next_=True, lines=True)
                if value.startswith("**Panel B."):
                    p.paragraph_format.page_break_before = True
        elif kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            add_inline(p, value)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, value)
        elif kind == "reference":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            add_inline(p, value, allow_italics=True)
            set_spacing(p, line=1.15, after=6)
        elif kind == "image":
            alt, rel_path = value
            path = ROOT / "Assignment Report" / rel_path
            last_image_dense = path.name in DENSE_FIGURE_ASSETS
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            if last_image_dense:
                p.paragraph_format.page_break_before = True
            set_spacing(p, line=1.0, before=4, after=2)
            run = p.add_run()
            run.add_picture(
                str(path),
                width=Inches(6.1 if last_image_dense else 5.7),
            )
            set_image_alt(run, FIGURE_ALT_TEXT.get(path.name, alt))
            set_keep(p, next_=True, lines=True)
            last_image = run
        elif kind in ("figure_caption", "table_caption"):
            p = doc.add_paragraph(style="Caption")
            add_inline(p, value)
            if value.startswith("Table 9. Reserved-test confirmation"):
                p.paragraph_format.page_break_before = True
            set_keep(p, next_=kind == "table_caption", lines=True)
            if kind == "figure_caption":
                last_image_dense = False
        elif kind == "table":
            rows = value
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_index, row in enumerate(rows):
                for c_index, text in enumerate(row):
                    cell = table.cell(r_index, c_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.right_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    add_inline(p, text)
                    set_spacing(p, line=1.15, after=0)
                    if r_index == 0:
                        for run in p.runs:
                            run.bold = True
                        shade_cell(cell)
            repeat_header(table.rows[0])
            for row_index, row in enumerate(table.rows):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        set_keep(
                            paragraph,
                            next_=row_index < len(table.rows) - 1,
                            lines=True,
                        )
            set_table_geometry(table, table_widths(rows))


def build(page_map):
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    title, abstract = get_front_text(md_text)
    blocks = parse_markdown(md_text)
    headings = [(kind, value) for kind, value in blocks if kind in ("h1", "h2")]
    figures = [value for kind, value in blocks if kind == "figure_caption"]
    tables = [value for kind, value in blocks if kind == "table_caption"]

    doc = Document(str(REFERENCE_DOCX))
    clear_document_body(doc)
    apply_styles(doc)
    first = doc.sections[0]
    configure_section(first)
    set_header(first, "", "")
    set_footer(first, enabled=False)
    first.different_first_page_header_footer = True
    first.first_page_header.paragraphs[0].clear()
    first.first_page_footer.paragraphs[0].clear()
    add_cover(doc)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front)
    front.different_first_page_header_footer = False
    set_header(front, "CT046-3-M-AML", "Applied Machine Learning")
    set_footer(front)
    set_page_number_format(front, "lowerRoman", start=1)
    add_front_matter(doc, title, abstract, headings, figures, tables, page_map)

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main)
    main.different_first_page_header_footer = False
    set_header(main, "CT046-3-M-AML", "GPA Change Prediction")
    set_footer(main)
    set_page_number_format(main, "decimal", start=1)
    add_body(doc, blocks)

    # Remove preserve-only layout tables inherited from the retained artifact;
    # the rebuilt headers and footers use ordinary paragraphs and tab stops.
    for section in doc.sections:
        for story in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for table in list(story.tables):
                table._element.getparent().remove(table._element)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is not None:
        settings.remove(update)
    do_not_update_links = settings.find(qn("w:doNotUpdateLinks"))
    if do_not_update_links is None:
        do_not_update_links = OxmlElement("w:doNotUpdateLinks")
        settings.append(do_not_update_links)
    do_not_update_links.set(qn("w:val"), "true")
    doc.core_properties.title = title
    doc.core_properties.subject = "Comparative machine learning study of semester GPA change"
    doc.core_properties.author = "Tanjim Noor"
    doc.save(str(OUTPUT_DOCX))
    mark_orphan_header_table_rows(OUTPUT_DOCX)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--page-map", type=Path)
    args = parser.parse_args()
    page_map = {}
    if args.page_map and args.page_map.exists():
        page_map = json.loads(args.page_map.read_text(encoding="utf-8"))
    build(page_map)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
